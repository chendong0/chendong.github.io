import docx
import jieba
import time
import os

# Start the timer
start_time = time.time()

# Define the base directory
base_dir = r"/home/chen/Downloads/DATA20231211/"

# Define input and output file paths using os.path.join
input_file_path = os.path.join(base_dir, "20240111docx.docx")
output_file_path = os.path.join(base_dir, "202312158clujieba.docx")

# 读取输入docx文档 
input_doc = docx.Document(input_file_path)

# 定义存储标注文本的列表
text = []

# 存储每段标注文本 
tagged_text = ""

# 词频统计字典
adj_counts = Counter()
adverb_counts = Counter()
noun_counts = Counter()
prep_counts = Counter()
verb_counts = Counter()

# 遍历每段文本进行处理  
for para in input_doc.paragraphs:
    
     # 使用jieba分词并标注词性
    words = jieba.posseg.cut(para.text)
    
    # 遍历分词结果 Update Counters based on POS Tags:
    for word, flag in words:
        
        # 拼接标注文本 
        if flag != 'x':
            tagged_text += f"{word}/{flag}"
            
           # 根据词性统计词频          
        if flag.startswith("a"):
            adj_counts[word] += 1 
        elif flag.startswith("a"):
            adverb_counts[word] += 1
        elif flag.startswith("n"):
            noun_counts[word] += 1
        elif flag.startswith("p"):
            prep_counts[word] += 1
        elif flag.startswith("v"):
            verb_counts[word] += 1
            
    # 将标注文本添加到列表         
    text.append(tagged_text)
    tagged_text = ""
    
# 输出到docx文档   
output_doc = docx.Document() 

# Output Tagged Text to Document:
for para in text:
    output_doc.add_paragraph(para)

    # Sorting Word Counts for Different Parts of Speech:
'''
lambda x: x[1] is an anonymous (lambda) function in Python.

x is the input argument (a tuple in this case).
x[1] extracts the second element of the tuple
在 Python 中，lambda 关键字可以用来创建匿名函数，即不需要使用 def 语句定义的函数。匿名函数通常用于简洁表达式，尤其是在其他函数或方法中。

在 sorted_adverb = sorted(adj_counts.items(), key=lambda x:x[1], reverse=True) 中，lambda x:x[1] 是一个匿名函数，它接受一个 x 参数，
并返回 x 的第二个元素。在本例中，x 是 adj_counts 字典的键值对。因此，lambda x:x[1] 实际上是说，对于每个键值对 x，返回其第二个元素，即计数。

sorted() 函数用于对列表或可迭代对象进行排序。在本例中，它用于对 adj_counts.items() 列表进行排序。key 参数指定排序的键。在本例中，我们指定使用 lambda x:x[1] 函数来确定排序顺序。因此，sorted_adverb 列表将按计数从高到低进行排序。
在 adj_counts 字典中，每个键值对都由两个元素组成：一个是形容词，另一个是形容词出现的次数。因此，x[1] 返回的第二个元素就是形容词出现的次数，即计数。

'''
sorted_adj = sorted(adj_counts.items(), key=lambda x:x[1], reverse=True)
sorted_adverb = sorted(adj_counts.items(), key=lambda x:x[1], reverse=True)
sorted_nouns = sorted(noun_counts.items(), key=lambda x:x[1], reverse=True)
sorted_preps = sorted(prep_counts.items(), key=lambda x:x[1], reverse=True)
sorted_verb = sorted(prep_counts.items(), key=lambda x:x[1], reverse=True)



"""
所有介词的词频,而不限制top 10,可以直接遍历整个排序后的介词词频列表,不指定切片:
output_doc.add_paragraph("\n介词词频:")
for prep, count in sorted_preps[10]:
    output_doc.add_paragraph(f"{prep}:{count}")
直接遍历 sorted_preps 这个已经按词频排序的列表,就可以输出所有的介词和词频,不再限制前10个。

"""

'''
如果想要限制输出词频最小值,可以添加if判断:
output_doc.add_paragraph("\n介词词频:")
for prep, count in sorted_preps:
    if count >= 5:
        output_doc.add_paragraph(f"{prep}:{count}")
'''

# Sort and Output Top 10 Nouns:

#Sort and Output Adjectives, Prepositions, and Verbs:
output_doc.add_paragraph("\n名词词频:")  
for noun,count in sorted_nouns[:10]:
    output_doc.add_paragraph(f"{noun}:{count}")
    
output_doc.add_paragraph("\nadj:")
for adj,count in sorted_adj:
    output_doc.add_paragraph(f"{adj}:{count}")
    
output_doc.add_paragraph("\npreps:")
for prep,count in sorted_preps:
    output_doc.add_paragraph(f"{prep}:{count}")
    
output_doc.add_paragraph("\nverb:")
for verb,count in sorted_verb:
    output_doc.add_paragraph(f"{verb}:{count}")

    #Save Output Document:
output_doc.save(output_file_path)

# End the timer
# End Timer and Display Execution Time:
end_time = time.time()
time_taken = end_time - start_time

print(f"输出文件已保存到:{output_file_path}")
print(time_taken)
