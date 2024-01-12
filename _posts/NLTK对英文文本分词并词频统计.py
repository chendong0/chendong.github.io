import nltk  # 导入自然语言处理工具包
from nltk.tokenize import word_tokenize  # 导入分词模块
from collections import Counter  # 导入计数器模块
import docx  # 导入文档处理模块
import time  # 导入时间模块用于计时

start_time = time.time()  # 记录程序开始时间

def process_paragraph(paragraph, noun_counts, adj_counts, verb_counts, prep_counts):
    """处理段落文本，统计词频并进行分词标注"""
    tokens = word_tokenize(paragraph)  # 分词
    tagged_text = ""  # 用来存储分词标注后的文本

    for word, flag in nltk.pos_tag(tokens):  # 获取每个词的词性标注
        if flag in ["NN", "VB", "JJ", "RB"]:  # 只处理指定词性的词语
            if flag.startswith("NN"):  # 名词
                noun_counts[word] += 1  # 统计名词词频
            elif flag.startswith("VB"):  # 动词
                verb_counts[word] += 1  # 统计动词词频
            elif flag.startswith("JJ"):  # 形容词
                adj_counts[word] += 1  # 统计形容词词频
            elif flag.startswith("RB"):  # 副词
                prep_counts[word] += 1  # 统计副词词频

            # 增加分词前后空格，便于识别
            tagged_text += f" {word} /{flag} "  # 添加分词标注
        else:
            # 非指定词性不添加分词标注
            tagged_text += f" {word}"

    return tagged_text.strip()  # 返回处理后的文本

def save_word_frequencies(doc, title, word_counts, limit=None):
    """保存词频到文档中"""
    doc.add_paragraph(f"\n{title}:")  # 添加标题
    sorted_words = sorted(word_counts.items(), key=lambda x: x[1], reverse=True)  # 按词频降序排序
    for word, count in sorted_words[:limit]:  # 限制输出词频前若干词
        doc.add_paragraph(f"{word}: {count}")  # 添加词频

def main(input_file_path, output_file_path):
    """主函数"""
    input_doc = docx.Document(input_file_path)  # 打开输入文档
    output_doc = docx.Document()  # 创建输出文档

    noun_counts = Counter()  # 名词词频计数器
    adj_counts = Counter()  # 形容词词频计数器
    verb_counts = Counter()  # 动词词频计数器
    prep_counts = Counter()  # 副词词频计数器

    for para in input_doc.paragraphs:  # 遍历输入文档的每个段落
        tagged_text = process_paragraph(para.text, noun_counts, adj_counts, verb_counts, prep_counts)  # 处理段落
        output_doc.add_paragraph(tagged_text)  # 将处理后的文本加入输出文档

    # 保存各词性的词频到输出文档中
    save_word_frequencies(output_doc, "NN名词词频", noun_counts)
    save_word_frequencies(output_doc, "JJ形容词词频", adj_counts)
    save_word_frequencies(output_doc, "VB动词词频", verb_counts)
    save_word_frequencies(output_doc, "RB副词词频", prep_counts)

    output_doc.save(output_file_path)  # 保存输出文档
    end_time = time.time()  # 记录程序结束时间
     
    print(f"输出文件已保存到: {output_file_path}")
    
if __name__ == "__main__":
    input_file_path = r"/home/chen/Downloads/DATA20231211/20240111docx.docx"
    output_file_path = r"/home/chen/Downloads/DATA20231211/202312150bard1.docx"
    main(input_file_path, output_file_path)

'''
词性分析的代码
Part-of-Speech Codes
CC Coordinating conjunction
CD Cardinal number
DT Determiner
EX Existential there
FW Foreign word
IN Preposition or subordinating 
conjunction
JJ Adjective
JJR Adjective, comparative
JJS Adjective, superlative
LS List item marker
MD Modal
NN Noun, singular or mass
NNS Noun, plural
NNP Proper noun, singular
NNPS Proper noun, plural
PDT Predeterminer
POS Possessive ending
PRP Personal pronoun
PRP$ Possessive pronoun
RB Adverb 
RBR Adverb, comparative
RBS Adverb, superlative
RP Particle
SYM Symbol
TO to
UH Interjection
VB Verb, base form
VBD Verb, past tense
VBG Verb, gerund or present 
participle
VBN Verb, past participle
VBP Verb, non-3rd person singular 
present
VBZ Verb, 3rd person singular 
present
WDT Wh-determiner
WP Wh-pronoun
WP$ Possessive wh-pronoun
WRB Wh-adver
'''
