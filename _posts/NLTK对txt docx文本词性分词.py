import nltk
from nltk.tokenize import word_tokenize
from collections import Counter
import docx
import time
import os

start_time = time.time()

def process_text(text, noun_counts, adj_counts, verb_counts, prep_counts):
    tokens = word_tokenize(text)
    tagged_text = ""

    for word, flag in nltk.pos_tag(tokens):
        if flag in ["NN", "VB", "JJ", "RB", "IN"]:  # 添加 "IN" 以支持介词词频统计
            if flag.startswith("NN"):
                noun_counts[word] += 1
            elif flag.startswith("VB"):
                verb_counts[word] += 1
            elif flag.startswith("JJ"):
                adj_counts[word] += 1
            elif flag.startswith("RB"):
                prep_counts[word] += 1  # 这里更适合用 "IN" 来统计介词频率
            elif flag.startswith("TO"):
                prep_counts[word] += 1  # 添加 "TO" 作为介词词性

            tagged_text += f" {word} /{flag} "
        else:
            tagged_text += f" {word}"

    return tagged_text.strip()

def main(input_file_path, output_file_path):
    input_doc = None
    output_doc = docx.Document()

    noun_counts = Counter()
    adj_counts = Counter()
    verb_counts = Counter()
    prep_counts = Counter()

    file_extension = os.path.splitext(input_file_path)[1].lower()
    if file_extension == ".docx":
        input_doc = docx.Document(input_file_path)
        for para in input_doc.paragraphs:
            tagged_text = process_text(para.text, noun_counts, adj_counts, verb_counts, prep_counts)
            output_doc.add_paragraph(tagged_text)
    elif file_extension in (".txt", ".doc"):
        with open(input_file_path, "r", encoding="utf-8") as input_file:
            text = input_file.read()
            tagged_text = process_text(text, noun_counts, adj_counts, verb_counts, prep_counts)
            output_doc.add_paragraph(tagged_text)
    else:
        print("不支持的文件格式")
        return

    save_word_frequencies(output_doc, "NN名词词频", noun_counts)
    save_word_frequencies(output_doc, "JJ形容词词频", adj_counts)
    save_word_frequencies(output_doc, "VB动词词频", verb_counts)
    save_word_frequencies(output_doc, "IN介词词频", prep_counts)  # 使用 "IN" 表示介词词频

    output_doc.save(output_file_path)
    end_time = time.time()
    time_taken = end_time - start_time

    print(f"输出文件已保存到: {output_file_path}")
    print(time_taken)

if __name__ == "__main__":
    input_file_path = r"/home/chen/Downloads/DATA20231211/20240111docx.docx"  # 可以替换为 .txt 或 .doc 文件
    output_file_path = r"/home/chen/Downloads/DATA20231211/202312150openai.docx"
    main(input_file_path, output_file_path)
