import nltk
from nltk.tokenize import word_tokenize
from collections import Counter
import docx

def process_paragraph(paragraph, noun_counts, adj_counts, verb_counts, prep_counts):
    """处理段落文本，统计指定词性的词频"""
    tokens = word_tokenize(paragraph)  # 对段落进行分词

    for word, flag in nltk.pos_tag(tokens):  # 获取每个词的词性标注
        if flag != 'TO':  # 过滤掉词性为 'TO' 的词（通常为虚词）
            if flag.startswith("NN"):  # 名词
                noun_counts[word] += 1  # 统计名词词频
            elif flag.startswith("VB"):  # 动词
                verb_counts[word] += 1  # 统计动词词频
            elif flag.startswith("JJ"):  # 形容词
                adj_counts[word] += 1  # 统计形容词词频
            elif flag.startswith("IN"):  # 介词
                prep_counts[word] += 1  # 统计介词词频

def save_word_frequencies(doc, title, word_counts, limit=10):
    """保存词频到文档中"""
    doc.add_paragraph(f"\n{title}:")  # 添加标题
    sorted_words = sorted(word_counts.items(), key=lambda x: x[1], reverse=True)  # 按词频降序排序
    for word, count in sorted_words[:limit]:  # 输出前 10 个词频最高的词
        doc.add_paragraph(f"{word}: {count}")  # 添加词频信息到文档

def main(input_file_path, output_file_path):
    """主函数"""
    input_doc = docx.Document(input_file_path)  # 打开输入文档
    output_doc = docx.Document()  # 创建输出文档

    noun_counts = Counter()  # 名词词频计数器
    adj_counts = Counter()  # 形容词词频计数器
    verb_counts = Counter()  # 动词词频计数器
    prep_counts = Counter()  # 介词词频计数器

    for para in input_doc.paragraphs:  # 遍历输入文档的每个段落
        process_paragraph(para.text, noun_counts, adj_counts, verb_counts, prep_counts)  # 处理段落

    # 保存各词性的词频到输出文档中
    save_word_frequencies(output_doc, "NN名词词频", noun_counts)
    save_word_frequencies(output_doc, "JJ形容词词频", adj_counts)
    save_word_frequencies(output_doc, "VB动词词频", verb_counts)
    save_word_frequencies(output_doc, "IN介词词频", prep_counts)

    output_doc.save(output_file_path)  # 保存输出文档
    print(f"输出文件已保存到: {output_file_path}")  # 输出保存路径

if __name__ == "__main__":
    # 设置输入输出文件路径（注意根据实际情况修改路径）
    input_file_path = r"/home/chen/Downloads/DATA20231211/20240111docx.docx"
    output_file_path = r"/home/chen/Downloads/DATA20231211/202312150bard1.docx"
    main(input_file_path, output_file_path)
